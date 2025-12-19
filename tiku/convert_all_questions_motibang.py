#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
多格式题库转换脚本 - 磨题帮版本
将题库文件转换为磨题帮导入格式

支持输入文件格式：
- .xlsx (Excel文件)
- .doc (旧版WPS文档)
- .docx (Word文档)

使用方法：
    python3 convert_all_questions_motibang.py [文件1] [文件2] ...

示例：
    # 转换单个文件
    python3 convert_all_questions_motibang.py 我的题库.xlsx

    # 转换多个文件
    python3 convert_all_questions_motibang.py 题库1.xlsx 题库2.docx

    # 转换当前目录所有Excel文件
    python3 convert_all_questions_motibang.py *.xlsx

输出格式：磨题帮Excel导入模板
输出文件：原文件名_磨题帮.xlsx

磨题帮模板格式（基于template1）：
- 题干
- 题型（选择题/判断题/填空题）
- 选择项1-10
- 答案
- 解析
- 得分
"""

import re
import os
import sys
import glob as glob_module
import argparse
import openpyxl
from openpyxl.styles import Font
from docx import Document
import olefile

# 版本信息
__version__ = '1.1.0'
__description__ = '多格式题库转换工具 - 磨题帮版本'


# ========== 字符规范化工具 ==========
def normalize_text(text):
    """
    全面的字符规范化：
    - 全角字母 → 半角字母
    - 全角数字 → 半角数字
    - 全角标点 → 半角标点（部分）
    """
    if not text:
        return ''

    result = []
    for char in text:
        code = ord(char)
        # 全角大写字母 A-Z: U+FF21 - U+FF3A
        if 0xFF21 <= code <= 0xFF3A:
            result.append(chr(code - 0xFF21 + ord('A')))
        # 全角小写字母 a-z: U+FF41 - U+FF5A
        elif 0xFF41 <= code <= 0xFF5A:
            result.append(chr(code - 0xFF41 + ord('a')))
        # 全角数字 0-9: U+FF10 - U+FF19
        elif 0xFF10 <= code <= 0xFF19:
            result.append(chr(code - 0xFF10 + ord('0')))
        # 全角空格
        elif code == 0x3000:
            result.append(' ')
        # 常用全角标点转换
        elif char == '．':
            result.append('.')
        elif char == '，':
            result.append(',')
        elif char == '：':
            result.append(':')
        elif char == '；':
            result.append(';')
        elif char == '（':
            result.append('(')
        elif char == '）':
            result.append(')')
        else:
            result.append(char)
    return ''.join(result)


def normalize_letters(text):
    """将全角字母转换为半角字母（保留向后兼容）"""
    return normalize_text(text)


def normalize_judgment_answer(raw_answer):
    """
    统一判断题答案格式，支持多种输入格式
    返回: '对' / '错' / '' (无法识别)
    """
    if not raw_answer:
        return ''

    ans = str(raw_answer).strip().upper()

    # 正确答案的各种表示
    true_values = [
        'Y', 'YES', 'T', 'TRUE', '1',
        '对', '正确', '√', '✓', '是', 'RIGHT', 'CORRECT',
        '○', '〇', 'O'
    ]

    # 错误答案的各种表示
    false_values = [
        'N', 'NO', 'F', 'FALSE', '0',
        '错', '错误', '×', '✗', '✘', '否', 'WRONG', 'INCORRECT',
        'X', '☓', 'Ⅹ',  # 包含罗马数字大写 X (U+2169)，因为 .upper() 会把 ⅹ 转成 Ⅹ
        '叉', '✕', '❌', '⨯', '⨉'  # 更多叉号变体
    ]

    if ans in true_values:
        return '对'
    elif ans in false_values:
        return '错'
    else:
        return ''


def clean_answer(answer):
    """清理答案字符串"""
    if not answer:
        return ''
    answer = str(answer).strip().upper()
    # 只保留字母
    return ''.join(c for c in answer if c.isalpha())


# ========== 处理 车辆检修工练习题-中级.xlsx ==========
def parse_mid_level_excel(file_path):
    """解析中级练习题Excel"""
    questions = []
    wb = openpyxl.load_workbook(file_path)
    ws = wb.active

    # 跳过表头行
    for row_idx in range(2, ws.max_row + 1):
        question_text = ws.cell(row=row_idx, column=2).value  # 试题内容
        if not question_text:
            continue

        # 获取题型（列16）
        q_type_raw = ws.cell(row=row_idx, column=16).value  # 题型列
        q_type_raw = str(q_type_raw).strip() if q_type_raw else ''

        # 获取答案（列15）
        answer_raw = ws.cell(row=row_idx, column=15).value
        answer_raw = str(answer_raw).strip() if answer_raw else ''

        # 根据题型处理
        if '判断' in q_type_raw:
            # 判断题：使用统一的答案规范化
            answer = normalize_judgment_answer(answer_raw)

            questions.append({
                'question': str(question_text).strip(),
                'type': '判断题',
                'answer': answer,
                'source': ''
            })
        else:
            # 选择题
            # 获取选项
            options = {}
            option_letters = 'ABCDEFGHIJKL'
            for i, letter in enumerate(option_letters):
                opt_value = ws.cell(row=row_idx, column=3+i).value
                if opt_value and str(opt_value).strip():
                    options[letter] = str(opt_value).strip()

            answer = clean_answer(answer_raw)

            # 根据答案数量判断题型
            if len(answer) > 1:
                q_type = '多选题'
            else:
                q_type = '单选题'

            questions.append({
                'question': str(question_text).strip(),
                'type': q_type,
                'options': options,
                'answer': answer,
                'source': ''
            })

    return questions


# ========== 处理 2-车辆题库汇总2024.xlsx ==========
def parse_2024_summary_excel(file_path):
    """解析2024题库汇总Excel"""
    questions = []
    wb = openpyxl.load_workbook(file_path)

    # 处理选择题工作表
    if '选择题' in wb.sheetnames:
        ws = wb['选择题']
        for row_idx in range(3, ws.max_row + 1):  # 从第3行开始（跳过标题行）
            question_text = ws.cell(row=row_idx, column=3).value  # 试题内容
            if not question_text:
                continue

            options = {}
            for i, letter in enumerate('ABCD'):
                opt_value = ws.cell(row=row_idx, column=4+i).value
                if opt_value and str(opt_value).strip():
                    # 选项可能带有字母前缀，需要去除
                    opt_text = str(opt_value).strip()
                    opt_text = re.sub(r'^[A-D][、.．:：]\s*', '', opt_text)
                    if opt_text:
                        options[letter] = opt_text

            answer = clean_answer(ws.cell(row=row_idx, column=8).value)

            questions.append({
                'question': str(question_text).strip(),
                'type': '单选题',
                'options': options,
                'answer': answer,
                'source': '',
                '_source_sheet': '选择题',
                '_source_row': row_idx
            })

    # 处理多选题工作表
    if '多选题' in wb.sheetnames:
        ws = wb['多选题']
        for row_idx in range(3, ws.max_row + 1):
            question_text = ws.cell(row=row_idx, column=3).value
            if not question_text:
                continue

            options = {}
            for i, letter in enumerate('ABCDE'):
                opt_value = ws.cell(row=row_idx, column=4+i).value
                if opt_value and str(opt_value).strip():
                    opt_text = str(opt_value).strip()
                    opt_text = re.sub(r'^[A-E][、.．:：]\s*', '', opt_text)
                    if opt_text:
                        options[letter] = opt_text

            answer = clean_answer(ws.cell(row=row_idx, column=9).value)

            questions.append({
                'question': str(question_text).strip(),
                'type': '多选题',
                'options': options,
                'answer': answer,
                'source': '',
                '_source_sheet': '多选题',
                '_source_row': row_idx
            })

    # 处理判断题工作表
    if '判断题' in wb.sheetnames:
        ws = wb['判断题']
        for row_idx in range(3, ws.max_row + 1):
            question_text = ws.cell(row=row_idx, column=3).value
            if not question_text:
                continue

            raw_answer = ws.cell(row=row_idx, column=4).value
            answer = normalize_judgment_answer(raw_answer)

            questions.append({
                'question': str(question_text).strip(),
                'type': '判断题',
                'answer': answer,
                'source': '',
                '_source_sheet': '判断题',
                '_source_row': row_idx
            })

    # 处理填空题工作表
    if '填空题' in wb.sheetnames:
        ws = wb['填空题']
        for row_idx in range(3, ws.max_row + 1):
            question_text = ws.cell(row=row_idx, column=3).value
            if not question_text:
                continue

            answer = ws.cell(row=row_idx, column=4).value or ''
            # 填空题答案可能用顿号、逗号分隔
            answers = re.split(r'[,，、;；]', str(answer))
            answers = [a.strip() for a in answers if a.strip()]

            questions.append({
                'question': str(question_text).strip(),
                'type': '填空题',
                'answers': answers,
                'raw_answer': str(answer),
                'source': '',
                '_source_sheet': '填空题',
                '_source_row': row_idx
            })

    # 处理简答题工作表
    if '简答题' in wb.sheetnames:
        ws = wb['简答题']
        for row_idx in range(3, ws.max_row + 1):
            question_text = ws.cell(row=row_idx, column=3).value
            if not question_text:
                continue

            answer = ws.cell(row=row_idx, column=4).value or ''

            questions.append({
                'question': str(question_text).strip(),
                'type': '简答题',
                'answer': str(answer).strip(),
                'source': '',
                '_source_sheet': '简答题',
                '_source_row': row_idx
            })

    # 处理论述题工作表
    if '论述题' in wb.sheetnames:
        ws = wb['论述题']
        for row_idx in range(3, ws.max_row + 1):
            question_text = ws.cell(row=row_idx, column=3).value
            if not question_text:
                continue

            answer = ws.cell(row=row_idx, column=4).value or ''

            questions.append({
                'question': str(question_text).strip(),
                'type': '简答题',  # 论述题归类为简答题
                'answer': str(answer).strip(),
                'source': '',
                '_source_sheet': '论述题',
                '_source_row': row_idx
            })

    return questions


# ========== 处理 题库1.doc ==========
def parse_doc_file(file_path):
    """解析旧版WPS文档"""
    questions = []

    try:
        ole = olefile.OleFileIO(file_path)
        wd = ole.openstream('WordDocument').read()
        text = wd.decode('utf-16-le', errors='ignore')

        # 提取中文文本块
        chinese_blocks = re.findall(
            r'[\u4e00-\u9fff\uff00-\uffef\u3000-\u303f0-9A-Za-z（）()、，。？！：；√×\s\-—_]+',
            text
        )

        # 合并文本
        full_text = ''.join(block for block in chinese_blocks if len(block) > 3)

        # 找到各部分的位置
        judgment_start = full_text.find('判断题')
        fill_start = full_text.find('填空题')
        essay_start = full_text.find('简答题')

        # ===== 解析选择题 =====
        if judgment_start > 0:
            choice_text = full_text[:judgment_start]
        else:
            choice_text = full_text

        # 按题号分割
        choice_splits = re.split(r'(?=\d+[、．.]\s*[^\d])', choice_text)

        for chunk in choice_splits:
            chunk = chunk.strip()
            if not chunk or len(chunk) < 20:
                continue

            # 提取答案
            answer_match = re.search(r'[（(]\s*([A-Da-d])\s*[)）]', chunk)
            if not answer_match:
                continue

            answer = answer_match.group(1).upper()

            # 提取题干（答案括号之前的部分，去掉题号）
            q_text_match = re.match(r'\d*[、．.]?\s*(.+?)[（(]\s*[A-Da-d]\s*[)）]', chunk, re.DOTALL)
            if not q_text_match:
                continue

            q_text = q_text_match.group(1).strip()
            q_text = re.sub(r'\s+', ' ', q_text)

            # 提取选项
            options = {}
            opt_text = chunk[answer_match.end():]

            for letter in 'ABCD':
                pattern = rf'{letter}[、．.]\s*([^A-D]+?)(?=[A-D][、．.]|$)'
                opt_match = re.search(pattern, opt_text)
                if opt_match:
                    opt_content = opt_match.group(1).strip()
                    opt_content = re.sub(r'\s+', ' ', opt_content)
                    if opt_content and len(opt_content) < 100:
                        options[letter] = opt_content

            if q_text and len(options) >= 2 and len(q_text) > 5:
                questions.append({
                    'question': q_text,
                    'type': '单选题',
                    'options': options,
                    'answer': answer,
                    'source': ''
                })

        # ===== 解析判断题 =====
        if judgment_start > 0:
            if fill_start > judgment_start:
                judgment_text = full_text[judgment_start:fill_start]
            elif essay_start > judgment_start:
                judgment_text = full_text[judgment_start:essay_start]
            else:
                judgment_text = full_text[judgment_start:]

            # 匹配判断题: 题号、题干。（√或×）
            judgment_pattern = r'(\d+)[、．.]\s*([^（(√×]+?)[。？！]?\s*[（(]\s*([√×])\s*[)）]'
            judgment_matches = re.findall(judgment_pattern, judgment_text)

            for match in judgment_matches:
                q_num, q_text, raw_answer = match
                answer = '对' if raw_answer == '√' else '错'
                q_text = q_text.strip()
                q_text = re.sub(r'\s+', ' ', q_text)
                if not q_text.endswith(('。', '？', '！')):
                    q_text += '。'

                if len(q_text) > 5:
                    questions.append({
                        'question': q_text,
                        'type': '判断题',
                        'answer': answer,
                        'source': ''
                    })

        # ===== 解析简答题 =====
        if essay_start > 0:
            essay_text = full_text[essay_start:]

            # 匹配简答题: 题号、题干（问号结尾）
            essay_pattern = r'(\d+)[、．.]\s*([^？?]+[？?])'
            essay_matches = re.findall(essay_pattern, essay_text)

            for match in essay_matches:
                q_num, q_text = match
                q_text = q_text.strip()
                q_text = re.sub(r'\s+', ' ', q_text)

                if q_text and len(q_text) > 10:
                    questions.append({
                        'question': q_text,
                        'type': '简答题',
                        'answer': '',  # 简答题答案复杂，暂不提取
                        'source': ''
                    })

        ole.close()

    except Exception as e:
        print(f'解析doc文件出错: {e}')
        import traceback
        traceback.print_exc()

    return questions


# ========== 处理 CRH6集团竞赛题库.docx ==========
def parse_fill_blank_question(text):
    """解析填空题"""
    match = re.match(r'^(\d+)[、.．]\s*', text)
    if match:
        text = text[match.end():]

    answer = ''
    question = text

    answer_match = re.search(r'答案[：:]\s*[（(]([^）)]+)[）)]', text)
    if answer_match:
        question = text[:answer_match.start()].strip()
        answer = answer_match.group(1).strip()
    else:
        answer_match2 = re.search(r'答案[：:]\s*[（(](.+?)$', text)
        if answer_match2:
            question = text[:answer_match2.start()].strip()
            answer = answer_match2.group(1).strip()
        else:
            answer_match3 = re.search(r'答案[：:]\s*(.+?)$', text)
            if answer_match3:
                question = text[:answer_match3.start()].strip()
                answer = answer_match3.group(1).strip()

    question = re.sub(r'_{2,}', '____', question)
    question = re.sub(r'\s{3,}', '____', question)

    answers = re.split(r'[,，、;；]', answer)
    answers = [a.strip() for a in answers if a.strip()]

    return {
        'question': question,
        'type': '填空题',
        'answers': answers,
        'raw_answer': answer
    }


def parse_choice_question(text):
    """解析选择/多选题"""
    text = normalize_letters(text)

    match = re.match(r'^(\d+)[、.．]\s*', text)
    if match:
        text = text[match.end():]

    answer = ''
    # 支持带逗号分隔的多选答案，如 (B,D) 或 (BD)
    answer_match = re.search(r'答案[：:]\s*[（(]\s*([A-Za-z,，\s]+)\s*[）)]', text)
    if answer_match:
        # 提取答案并移除逗号、空格，只保留字母
        raw_answer = answer_match.group(1)
        answer = ''.join(c.upper() for c in raw_answer if c.isalpha())
        text_for_options = text[:answer_match.start()] + text[answer_match.end():]
    else:
        text_for_options = text

    options = {}
    lines = text_for_options.replace('\t', '  ').split('\n')
    all_text = ' '.join(lines)

    # 改进版正则，正确处理短选项
    option_pattern = r'([A-L])[、.．:：]\s*([^\s]+(?:\s+[^\sA-L][^\s]*)*)'
    matches = re.findall(option_pattern, all_text)

    for letter, content in matches:
        content = content.strip()
        content = re.sub(r'\s+$', '', content)
        if content:
            options[letter.upper()] = content

    if len(options) < 2:
        for part in re.split(r'\s{2,}|\t|\n', text_for_options):
            part = part.strip()
            opt_match = re.match(r'([A-L])[、.．:：]\s*(.+)', part)
            if opt_match:
                letter = opt_match.group(1).upper()
                content = opt_match.group(2).strip()
                if content:
                    options[letter] = content

    first_option = re.search(r'[A-L][、.．:：]', text_for_options)
    if first_option:
        question = text_for_options[:first_option.start()].strip()
    else:
        question = text_for_options.strip()

    question = question.rstrip()
    # 清理题干末尾可能残留的答案标记（支持带逗号的格式）
    question = re.sub(r'答案[：:]\s*[（(]?[A-Za-z,，\s]*[）)]?\s*$', '', question).strip()
    # 清理可能残留的 ",X）" 格式
    question = re.sub(r'[,，]\s*[A-Za-z]\s*[）)]\s*$', '', question).strip()

    question_type = '多选题' if len(answer) > 1 else '单选题'

    return {
        'question': question,
        'type': question_type,
        'options': options,
        'answer': answer
    }


def parse_judgment_question(text):
    """解析判断题"""
    match = re.match(r'^(\d+)[、.．]\s*', text)
    if match:
        text = text[match.end():]

    answer_match = re.search(r'答案[：:]\s*[（(]([√×对错]|正确|错误)[）)]', text)
    if answer_match:
        raw_answer = answer_match.group(1).strip()
        question = text[:answer_match.start()].strip()

        if raw_answer in ['√', '对', '正确']:
            answer = '对'
        else:
            answer = '错'
    else:
        question = text.strip()
        answer = ''

    return {
        'question': question,
        'type': '判断题',
        'answer': answer
    }


def parse_essay_question(text):
    """解析简答题"""
    match = re.match(r'^(\d+)[、.．]\s*', text)
    if match:
        text = text[match.end():]

    question = ''
    answer = ''

    lines = text.strip().split('\n')

    if len(lines) >= 2:
        question_lines = []
        answer_start = 0
        for i, line in enumerate(lines):
            question_lines.append(line)
            if '？' in line or '?' in line:
                answer_start = i + 1
                break

        if answer_start > 0 and answer_start < len(lines):
            question = '\n'.join(question_lines).strip()
            answer = '\n'.join(lines[answer_start:]).strip()
        else:
            question_end = -1
            for i, char in enumerate(text):
                if char in '？?':
                    question_end = i + 1
                    break

            if question_end > 0:
                question = text[:question_end].strip()
                answer = text[question_end:].strip()
            else:
                question = lines[0].strip()
                answer = '\n'.join(lines[1:]).strip() if len(lines) > 1 else ''
    else:
        question_end = -1
        for i, char in enumerate(text):
            if char in '？?':
                question_end = i + 1
                break

        if question_end > 0 and question_end < len(text):
            question = text[:question_end].strip()
            answer = text[question_end:].strip()
        else:
            question = text.strip()
            answer = ''

    return {
        'question': question,
        'type': '简答题',
        'answer': answer
    }


def parse_crh6_docx(file_path):
    """解析CRH6集团竞赛题库.docx"""
    questions = []
    doc = Document(file_path)

    # 处理表格0 - 填空题和选择题
    if len(doc.tables) > 0:
        table0 = doc.tables[0]
        for row in table0.rows:
            text = row.cells[0].text.strip()
            if not text:
                continue

            has_option = any(marker in text for marker in ['A、', 'A.', 'A．', 'A：', 'A:', '\nA', '\tA'])

            if has_option:
                q = parse_choice_question(text)
            else:
                q = parse_fill_blank_question(text)

            questions.append(q)

    # 处理表格1 - 判断题
    if len(doc.tables) > 1:
        table1 = doc.tables[1]
        for row in table1.rows:
            text = row.cells[0].text.strip()
            if not text:
                continue

            q = parse_judgment_question(text)
            questions.append(q)

    # 处理表格2 - 简答题
    if len(doc.tables) > 2:
        table2 = doc.tables[2]
        for row in table2.rows:
            text = row.cells[0].text.strip()
            if not text:
                continue

            q = parse_essay_question(text)
            questions.append(q)

    return questions


# ========== 通用纯文本解析器 ==========
def detect_encoding(file_path, sample_size=8192):
    """
    检测文本文件编码（优化版：只读取文件头部）

    Args:
        file_path: 文件路径
        sample_size: 采样字节数，默认8KB足够检测大多数编码

    Returns:
        检测到的编码名称
    """
    encodings = ['utf-8', 'gbk', 'gb2312', 'gb18030', 'utf-16', 'big5']

    # 读取文件头部用于编码检测
    try:
        with open(file_path, 'rb') as f:
            raw_data = f.read(sample_size)
    except IOError:
        return 'utf-8'

    for enc in encodings:
        try:
            raw_data.decode(enc)
            return enc
        except (UnicodeDecodeError, UnicodeError):
            continue
    return 'utf-8'


def parse_text_question(text):
    """
    智能解析单个题目文本，自动检测题型
    支持格式:
    - 选择题: 带A、B、C、D选项
    - 判断题: 答案为√/×/对/错/T/F等
    - 填空题: 带____或空格占位符
    - 简答题: 其他
    """
    text = normalize_text(text.strip())
    if not text or len(text) < 5:
        return None

    # 去除题号
    text = re.sub(r'^\d+[、.．:：\s]+', '', text)

    # 检测是否有选项标志 → 选择题
    has_options = bool(re.search(r'[A-L][、.．:：]\s*\S', text))

    if has_options:
        return parse_choice_question(text)

    # 检测是否有判断题答案标志
    judgment_pattern = r'答案[：:]\s*[（(]?\s*([√×对错TFYNtfyn]|正确|错误|TRUE|FALSE|Yes|No)\s*[）)]?'
    judgment_match = re.search(judgment_pattern, text, re.IGNORECASE)
    if judgment_match:
        question = text[:judgment_match.start()].strip()
        answer = normalize_judgment_answer(judgment_match.group(1))
        return {
            'question': question,
            'type': '判断题',
            'answer': answer
        }

    # 检测是否有填空标志
    has_blank = bool(re.search(r'_{2,}|（\s*）|\(\s*\)|【\s*】', text))
    if has_blank:
        return parse_fill_blank_question(text)

    # 检测是否有问号 → 可能是简答题
    if '？' in text or '?' in text:
        return parse_essay_question(text)

    # 默认作为简答题处理
    return {
        'question': text,
        'type': '简答题',
        'answer': ''
    }


def parse_text_file(file_path):
    """
    解析纯文本文件（.txt）
    支持多种题目分隔格式:
    - 数字编号: 1. / 1、/ 1) / (1)
    - 空行分隔
    - 特殊标记: 【题目】等
    """
    questions = []

    encoding = detect_encoding(file_path)
    try:
        with open(file_path, 'r', encoding=encoding) as f:
            content = f.read()
    except Exception as e:
        print(f'读取文本文件出错: {e}')
        return questions

    # 规范化内容
    content = normalize_text(content)

    # 尝试按题号分割
    # 匹配: 1. / 1、/ 1) / (1) / 【1】等格式
    split_pattern = r'\n\s*(?=\d+[、.．:：\)\)]\s|\(\d+\)|\【\d+\】)'
    chunks = re.split(split_pattern, content)

    if len(chunks) < 2:
        # 如果按题号分割效果不好，尝试按空行分割
        chunks = re.split(r'\n\s*\n+', content)

    for chunk in chunks:
        chunk = chunk.strip()
        if not chunk or len(chunk) < 10:
            continue

        q = parse_text_question(chunk)
        if q and q.get('question'):
            questions.append(q)

    return questions


def parse_generic_excel(file_path):
    """
    通用Excel解析器
    尝试自动检测Excel文件结构并解析题目
    """
    questions = []

    try:
        wb = openpyxl.load_workbook(file_path)
    except Exception as e:
        print(f'读取Excel文件出错: {e}')
        return questions

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]

        # 跳过空表
        if ws.max_row < 2:
            continue

        # 尝试检测表头行
        header_row = 1
        for row_idx in range(1, min(5, ws.max_row + 1)):
            cell_val = ws.cell(row=row_idx, column=1).value
            if cell_val and ('题' in str(cell_val) or '问题' in str(cell_val)):
                header_row = row_idx
                break

        # 遍历数据行
        for row_idx in range(header_row + 1, ws.max_row + 1):
            # 尝试从第一列或第二列获取题干
            question_text = None
            for col in [1, 2, 3]:
                val = ws.cell(row=row_idx, column=col).value
                if val and len(str(val).strip()) > 5:
                    question_text = str(val).strip()
                    break

            if not question_text:
                continue

            # 尝试组合所有单元格内容
            row_content = []
            for col in range(1, min(ws.max_column + 1, 20)):
                val = ws.cell(row=row_idx, column=col).value
                if val:
                    row_content.append(str(val).strip())

            full_text = ' '.join(row_content)
            q = parse_text_question(full_text)
            if q and q.get('question'):
                questions.append(q)

    return questions


def parse_generic_docx(file_path):
    """
    通用Word文档解析器
    尝试解析各种格式的docx文件
    """
    questions = []

    try:
        doc = Document(file_path)
    except Exception as e:
        print(f'读取Word文档出错: {e}')
        return questions

    # 首先尝试从表格中提取
    for table in doc.tables:
        for row in table.rows:
            text = ''
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    text += cell_text + ' '
            text = text.strip()
            if text and len(text) > 10:
                q = parse_text_question(text)
                if q and q.get('question'):
                    questions.append(q)

    # 如果表格没有内容，从段落中提取
    if not questions:
        all_text = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                all_text.append(text)

        full_content = '\n'.join(all_text)

        # 按题号分割
        chunks = re.split(r'\n\s*(?=\d+[、.．:：\)\)]\s)', full_content)
        for chunk in chunks:
            chunk = chunk.strip()
            if chunk and len(chunk) > 10:
                q = parse_text_question(chunk)
                if q and q.get('question'):
                    questions.append(q)

    return questions


def get_parser_for_file(file_path):
    """
    根据文件扩展名返回合适的解析器
    优先返回特定解析器，否则返回通用解析器
    """
    ext = os.path.splitext(file_path)[1].lower()
    base_name = os.path.basename(file_path).lower()

    # 特定文件使用特定解析器
    if '中级' in base_name and ext == '.xlsx':
        return parse_mid_level_excel
    elif '2024' in base_name and '汇总' in base_name and ext == '.xlsx':
        return parse_2024_summary_excel
    elif 'crh6' in base_name and ext == '.docx':
        return parse_crh6_docx
    elif ext == '.doc':
        return parse_doc_file

    # 通用解析器
    if ext == '.xlsx':
        return parse_generic_excel
    elif ext == '.docx':
        return parse_generic_docx
    elif ext == '.txt':
        return parse_text_file
    elif ext == '.doc':
        return parse_doc_file

    return None


# ========== 生成磨题帮Excel输出 ==========
def convert_to_motibang_excel(questions, output_path):
    """将题目转换为磨题帮Excel格式"""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = '题库'

    # 磨题帮表头（基于template1格式）
    headers = ['题干', '题型', '选择项1', '选择项2', '选择项3', '选择项4', '选择项5',
               '选择项6', '选择项7', '选择项8', '选择项9', '选择项10', '答案', '解析', '得分']

    # 添加表头
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)

    # 添加题目数据
    for row_idx, q in enumerate(questions, 2):
        q_type = q.get('type', '')

        # 题干
        ws.cell(row=row_idx, column=1, value=q.get('question', ''))

        # 题型转换为磨题帮格式
        if q_type in ['单选题', '多选题']:
            ws.cell(row=row_idx, column=2, value='选择题')
        elif q_type == '判断题':
            ws.cell(row=row_idx, column=2, value='判断题')
        elif q_type in ['填空题', '定序填空题', '不定序填空题']:
            ws.cell(row=row_idx, column=2, value='填空题')
        elif q_type == '简答题':
            ws.cell(row=row_idx, column=2, value='简答题')
        else:
            ws.cell(row=row_idx, column=2, value=q_type)

        if q_type in ['单选题', '多选题']:
            # 选择题选项
            options = q.get('options', {})
            sorted_options = sorted([(k, v) for k, v in options.items() if v], key=lambda x: x[0])

            # 写入选项（从列3开始）
            for i, (_, content) in enumerate(sorted_options[:10]):
                ws.cell(row=row_idx, column=3+i, value=content)

            # 答案映射
            original_answer = q.get('answer', '')
            if original_answer:
                old_letters = [k for k, _ in sorted_options]
                option_letters = 'ABCDEFGHIJ'
                old_to_new = {}
                for i, old_letter in enumerate(old_letters):
                    if i < len(option_letters):
                        old_to_new[old_letter] = option_letters[i]

                new_answer = ''
                for char in original_answer:
                    if char in old_to_new:
                        new_answer += old_to_new[char]
                    elif char.upper() in old_to_new:
                        new_answer += old_to_new[char.upper()]

                ws.cell(row=row_idx, column=13, value=new_answer if new_answer else original_answer)

        elif q_type == '判断题':
            # 判断题答案直接写入
            ws.cell(row=row_idx, column=13, value=q.get('answer', ''))

        elif q_type in ['填空题', '定序填空题']:
            # 填空题：答案用||分隔
            answers = q.get('answers', [])
            raw_answer = q.get('raw_answer', '')
            if answers:
                ws.cell(row=row_idx, column=13, value='||'.join(answers))
            elif raw_answer:
                # 将各种分隔符统一为||
                unified = re.sub(r'[,，、;；]', '||', raw_answer)
                ws.cell(row=row_idx, column=13, value=unified)

        elif q_type == '简答题':
            # 简答题答案
            ws.cell(row=row_idx, column=13, value=q.get('answer', ''))

    # 调整列宽
    ws.column_dimensions['A'].width = 60  # 题干
    ws.column_dimensions['B'].width = 10  # 题型
    for i, col in enumerate('CDEFGHIJKL'):
        ws.column_dimensions[col].width = 15  # 选项
    ws.column_dimensions['M'].width = 15  # 答案
    ws.column_dimensions['N'].width = 30  # 解析
    ws.column_dimensions['O'].width = 8   # 得分

    wb.save(output_path)
    print(f'已保存到: {output_path}')
    return len(questions)


def print_statistics(questions, title):
    """打印题目统计信息"""
    type_count = {}
    for q in questions:
        t = q.get('type', '未知')
        type_count[t] = type_count.get(t, 0) + 1

    print(f'\n{title} 题型统计:')
    for t, c in sorted(type_count.items()):
        print(f'  - {t}: {c}')


def deduplicate_questions(questions):
    """
    题目去重：基于题干+选项+答案综合判断，保留第一次出现的题目
    注意：题干相同但选项或答案不同的视为不同题目
    返回: (去重后的题目列表, 移除的重复数量, 重复题目详情列表)
    重复题目详情格式: [(重复题位置, 首次出现位置, 题干摘要, 题型), ...]
    位置格式: "工作表名 第N行" 或 "第N题"（无原始位置时）
    """
    seen = {}  # key -> (首次出现的题目对象, 列表索引)
    unique_questions = []
    duplicates_removed = 0
    duplicate_details = []  # 记录重复题目详情

    for idx, q in enumerate(questions):
        # 构建唯一标识：题干 + 选项内容 + 答案
        question_text = q.get('question', '').strip()
        if not question_text:
            continue

        # 选择题：包含选项内容
        options = q.get('options', {})
        if options:
            # 按字母顺序拼接选项内容
            options_str = '|'.join(f"{k}:{v}" for k, v in sorted(options.items()))
        else:
            options_str = ''

        # 获取答案（选择题/判断题用answer，填空题用answers）
        answer = q.get('answer', '')
        answers = q.get('answers', [])
        if answers:
            answer_str = '||'.join(answers)
        else:
            answer_str = str(answer)

        # 组合唯一键
        unique_key = f"{question_text}###{options_str}###{answer_str}"

        if unique_key not in seen:
            seen[unique_key] = (q, idx + 1)  # 存储题目对象和列表索引
            unique_questions.append(q)
        else:
            duplicates_removed += 1
            first_q, first_list_idx = seen[unique_key]

            # 获取当前重复题的原始位置
            dup_sheet = q.get('_source_sheet', '')
            dup_row = q.get('_source_row', 0)
            if dup_sheet and dup_row:
                dup_location = f"[{dup_sheet}] 第{dup_row}行"
            else:
                dup_location = f"第{idx + 1}题"

            # 获取首次出现题的原始位置
            first_sheet = first_q.get('_source_sheet', '')
            first_row = first_q.get('_source_row', 0)
            if first_sheet and first_row:
                first_location = f"[{first_sheet}] 第{first_row}行"
            else:
                first_location = f"第{first_list_idx}题"

            q_type = q.get('type', '未知')
            q_preview = question_text[:40] + ('...' if len(question_text) > 40 else '')
            duplicate_details.append((dup_location, first_location, q_preview, q_type))

    return unique_questions, duplicates_removed, duplicate_details


def validate_question(q, idx=0):
    """
    验证单个题目的完整性和有效性

    Args:
        q: 题目字典
        idx: 题目序号（用于输出）

    Returns:
        (是否有效, 警告列表)
    """
    warnings = []
    q_type = q.get('type', '')

    # 题干验证
    question_text = q.get('question', '').strip()
    if not question_text:
        return False, ['题干为空']
    if len(question_text) < 5:
        warnings.append('题干过短（<5字符）')

    # 选择题验证
    if q_type in ['单选题', '多选题']:
        options = q.get('options', {})
        if len(options) < 2:
            warnings.append(f'选项不足（仅{len(options)}个）')

        answer = q.get('answer', '')
        if not answer:
            warnings.append('答案为空')
        else:
            for char in answer:
                if char not in options:
                    warnings.append(f'答案"{char}"不在选项中')
                    break

    # 判断题验证
    elif q_type == '判断题':
        answer = q.get('answer', '')
        if answer and answer not in ['对', '错']:
            warnings.append(f'判断题答案格式异常: {answer}')

    # 填空题验证
    elif q_type in ['填空题', '定序填空题', '不定序填空题']:
        answers = q.get('answers', [])
        raw_answer = q.get('raw_answer', '')
        if not answers and not raw_answer:
            warnings.append('填空题答案为空')

    return True, warnings


def validate_questions(questions, verbose=False):
    """
    批量验证题目，返回验证统计

    Args:
        questions: 题目列表
        verbose: 是否输出每个警告

    Returns:
        (有效题目数, 警告总数, 警告详情列表)
    """
    valid_count = 0
    total_warnings = 0
    warning_details = []

    for idx, q in enumerate(questions, 1):
        valid, warnings = validate_question(q, idx)
        if valid:
            valid_count += 1
        if warnings:
            total_warnings += len(warnings)
            warning_details.append((idx, q.get('question', '')[:30], warnings))
            if verbose:
                for w in warnings:
                    print(f'  ⚠ 第{idx}题: {w}')

    return valid_count, total_warnings, warning_details


def get_file_base_name(file_path):
    """获取文件基础名称（不含扩展名）"""
    return os.path.splitext(os.path.basename(file_path))[0]


def process_file(file_path, parser_func, base_path, suffix='_磨题帮', verbose=False, dry_run=False):
    """
    通用文件处理函数

    Args:
        file_path: 输入文件路径
        parser_func: 解析器函数
        base_path: 输出目录
        suffix: 输出文件后缀
        verbose: 是否显示详细信息（包括验证警告）
        dry_run: 仅解析验证，不输出文件

    Returns:
        (questions, output_file, warnings_count) 或 (None, None, 0) 如果失败
    """
    if not os.path.exists(file_path):
        return None, None, 0

    file_name = os.path.basename(file_path)
    base_name = get_file_base_name(file_path)

    print('\n' + '-' * 40)
    print(f'处理 {file_name}...')

    # 解析文件（带异常处理）
    try:
        questions = parser_func(file_path)
    except Exception as e:
        print(f'  ✗ 解析失败: {e}')
        return None, None, 0

    if not questions:
        print(f'  ✗ 未解析到任何题目')
        return None, None, 0

    # 强制使用文件基础名作为source
    for q in questions:
        q['source'] = base_name

    print(f'  解析到 {len(questions)} 道题目')

    # 验证题目
    valid_count, warnings_count, warning_details = validate_questions(questions, verbose=verbose)
    if warnings_count > 0:
        print(f'  ⚠ 发现 {warnings_count} 个数据质量警告')
        if verbose and warning_details:
            for idx, q_text, warns in warning_details[:10]:  # 最多显示10个
                print(f'    第{idx}题 "{q_text}...": {", ".join(warns)}')
            if len(warning_details) > 10:
                print(f'    ... 还有 {len(warning_details) - 10} 个警告')

    # 去重处理
    questions, duplicates_removed, duplicate_details = deduplicate_questions(questions)
    if duplicates_removed > 0:
        print(f'  去除重复题目: {duplicates_removed} 道，剩余 {len(questions)} 道')
        # 输出重复题目详情
        for dup_loc, first_loc, q_preview, q_type in duplicate_details:
            print(f'    ↳ {dup_loc} [{q_type}] 与 {first_loc} 重复: "{q_preview}"')

    print_statistics(questions, base_name)

    # 如果是 dry-run 模式，不输出文件
    if dry_run:
        print(f'  [dry-run] 跳过文件输出')
        return questions, None, warnings_count

    output_name = f'{base_name}{suffix}.xlsx'
    output_path = os.path.join(base_path, output_name)
    convert_to_motibang_excel(questions, output_path)

    return questions, output_name, warnings_count


def create_argument_parser():
    """创建命令行参数解析器"""
    parser = argparse.ArgumentParser(
        prog='convert_all_questions_motibang.py',
        description=__description__,
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
    # 转换单个文件
    python3 %(prog)s 我的题库.xlsx

    # 转换多个文件
    python3 %(prog)s 题库1.xlsx 题库2.docx

    # 转换当前目录所有Excel文件
    python3 %(prog)s *.xlsx

    # 指定输出目录
    python3 %(prog)s -o /path/to/output 题库.xlsx

    # 详细模式（显示验证警告）
    python3 %(prog)s -v 题库.xlsx

    # 仅验证不输出文件
    python3 %(prog)s --dry-run 题库.xlsx

支持的文件格式: .xlsx, .docx, .doc, .txt
输出: 原文件名_磨题帮.xlsx
"""
    )
    parser.add_argument('files', nargs='*', metavar='文件',
                        help='要转换的题库文件（支持通配符）')
    parser.add_argument('-o', '--output', metavar='目录',
                        help='指定输出目录（默认为文件所在目录）')
    parser.add_argument('-v', '--verbose', action='store_true',
                        help='详细模式，显示验证警告信息')
    parser.add_argument('--dry-run', action='store_true',
                        help='仅解析验证，不输出文件')
    parser.add_argument('--version', action='version',
                        version=f'%(prog)s {__version__}')
    return parser


def main():
    """主函数，支持命令行参数"""
    # 解析命令行参数
    parser = create_argument_parser()
    args = parser.parse_args()

    # 展开通配符
    files_to_process = []
    for pattern in args.files:
        if '*' in pattern or '?' in pattern:
            expanded = glob_module.glob(pattern)
            files_to_process.extend(expanded)
        else:
            files_to_process.append(pattern)

    print('=' * 60)
    print(f'{__description__} v{__version__}')
    print('=' * 60)

    # 如果没有指定文件，显示帮助信息并退出
    if not files_to_process:
        print('错误: 未指定要转换的题库文件！\n')
        parser.print_help()
        return

    # 统计变量
    total_questions = 0
    total_warnings = 0
    generated_files = []
    failed_files = []
    skipped_files = []

    # 处理命令行指定的文件
    print(f'待处理文件: {len(files_to_process)} 个')
    if args.verbose:
        print(f'模式: {"仅验证" if args.dry_run else "转换输出"}')

    for file_path in files_to_process:
        # 检查文件是否存在
        if not os.path.exists(file_path):
            print(f'警告: 文件不存在，跳过 - {file_path}')
            skipped_files.append((file_path, '文件不存在'))
            continue

        # 检查文件扩展名
        ext = os.path.splitext(file_path)[1].lower()
        supported_exts = ['.xlsx', '.docx', '.doc', '.txt']
        if ext not in supported_exts:
            print(f'警告: 不支持的文件格式，跳过 - {file_path}')
            skipped_files.append((file_path, '不支持的格式'))
            continue

        # 获取解析器
        parser_func = get_parser_for_file(file_path)
        if not parser_func:
            print(f'警告: 无法找到合适的解析器，跳过 - {file_path}')
            skipped_files.append((file_path, '无合适解析器'))
            continue

        # 确定输出目录
        out_dir = args.output if args.output else os.path.dirname(os.path.abspath(file_path))
        if not out_dir:
            out_dir = '.'

        # 处理文件
        questions, output_name, warnings_count = process_file(
            file_path, parser_func, out_dir,
            verbose=args.verbose, dry_run=args.dry_run
        )

        if questions:
            if output_name:
                generated_files.append(output_name)
            total_questions += len(questions)
            total_warnings += warnings_count
        else:
            failed_files.append((file_path, '解析失败或无题目'))

    # 总结报告
    print('\n' + '=' * 60)
    print('转换报告')
    print('=' * 60)
    print(f'处理题目总数: {total_questions} 道')

    if generated_files:
        print(f'\n✓ 成功生成文件 ({len(generated_files)} 个):')
        for f in generated_files:
            print(f'  - {f}')

    if total_warnings > 0:
        print(f'\n⚠ 数据质量警告: {total_warnings} 个')
        if not args.verbose:
            print('  (使用 -v 参数查看详细警告信息)')

    if failed_files:
        print(f'\n✗ 处理失败 ({len(failed_files)} 个):')
        for f, reason in failed_files:
            print(f'  - {f}: {reason}')

    if skipped_files:
        print(f'\n○ 已跳过 ({len(skipped_files)} 个):')
        for f, reason in skipped_files:
            print(f'  - {f}: {reason}')

    if not generated_files and not args.dry_run:
        print('\n未生成任何文件！')

    print('=' * 60)


if __name__ == '__main__':
    main()
